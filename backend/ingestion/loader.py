"""
backend/ingestion/loader.py
----------------------------
Document loaders for PDF, DOCX, and plain text files.

Reads raw files and returns LangChain Document objects
(each with page_content and metadata).
"""

import logging
from pathlib import Path
from typing import List

from langchain_core.documents import Document

logger = logging.getLogger(__name__)


def load_pdf(file_path: str) -> List[Document]:
    """
    Load a PDF file and return one Document per page.
    Falls back to pypdf if unstructured fails.
    """
    path = Path(file_path)
    docs: List[Document] = []

    try:
        from unstructured.partition.pdf import partition_pdf
        elements = partition_pdf(filename=str(path))
        text = "\n\n".join([str(e) for e in elements if str(e).strip()])
        docs.append(Document(
            page_content=text,
            metadata={"source": path.name, "type": "pdf", "loader": "unstructured"},
        ))
        logger.info(f"Loaded PDF via unstructured: {path.name} ({len(text)} chars)")
    except Exception as e:
        logger.warning(f"unstructured failed ({e}), falling back to pypdf")
        from pypdf import PdfReader
        reader = PdfReader(str(path))
        for i, page in enumerate(reader.pages):
            text = page.extract_text() or ""
            if text.strip():
                docs.append(Document(
                    page_content=text,
                    metadata={"source": path.name, "type": "pdf", "page": i + 1, "loader": "pypdf"},
                ))
        logger.info(f"Loaded PDF via pypdf: {path.name} ({len(reader.pages)} pages)")

    return docs


def load_docx(file_path: str) -> List[Document]:
    """
    Load a DOCX file and return a single Document with all paragraph text.
    """
    path = Path(file_path)
    from docx import Document as DocxDocument

    doc = DocxDocument(str(path))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    text = "\n\n".join(paragraphs)

    logger.info(f"Loaded DOCX: {path.name} ({len(paragraphs)} paragraphs)")
    return [Document(
        page_content=text,
        metadata={"source": path.name, "type": "docx", "paragraphs": len(paragraphs)},
    )]


def load_text(file_path: str) -> List[Document]:
    """Load a plain .txt file."""
    path = Path(file_path)
    text = path.read_text(encoding="utf-8", errors="ignore")
    logger.info(f"Loaded TXT: {path.name} ({len(text)} chars)")
    return [Document(
        page_content=text,
        metadata={"source": path.name, "type": "txt"},
    )]


def load_document(file_path: str) -> List[Document]:
    """
    Auto-detect file type and load accordingly.

    Args:
        file_path: Absolute or relative path to the document.

    Returns:
        List of LangChain Document objects.

    Raises:
        ValueError: If file type is unsupported.
    """
    path = Path(file_path)
    suffix = path.suffix.lower()

    loaders = {
        ".pdf": load_pdf,
        ".docx": load_docx,
        ".doc": load_docx,
        ".txt": load_text,
    }

    if suffix not in loaders:
        raise ValueError(f"Unsupported file type: '{suffix}'. Supported: {list(loaders.keys())}")

    return loaders[suffix](file_path)
