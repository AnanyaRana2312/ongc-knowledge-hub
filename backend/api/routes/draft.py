"""
backend/api/routes/draft.py
---------------------------
Drafts an automated report based on a prompt and domain, and returns it as a DOCX or PDF file.
Uses the same RAG chain as the chat endpoint to avoid VRAM issues.
"""

import logging
import os
import uuid
import tempfile
from typing import Optional

from fastapi import APIRouter, HTTPException
from fastapi.responses import FileResponse
from pydantic import BaseModel, Field

logger = logging.getLogger(__name__)

router = APIRouter(prefix="/draft", tags=["Draft"])


class DraftRequest(BaseModel):
    topic: str = Field(..., description="The topic or prompt for the draft.")
    domain: Optional[str] = Field(None, description="Optional domain constraint.")


def generate_docx(markdown_text: str, filepath: str):
    try:
        from docx import Document
    except ImportError:
        raise RuntimeError("python-docx is not installed")
    
    doc = Document()
    
    for line in markdown_text.split('\n'):
        if line.startswith('# '):
            doc.add_heading(line.lstrip('# ').strip(), level=1)
        elif line.startswith('## '):
            doc.add_heading(line.lstrip('## ').strip(), level=2)
        elif line.startswith('### '):
            doc.add_heading(line.lstrip('### ').strip(), level=3)
        elif line.strip():
            doc.add_paragraph(line.strip())
            
    doc.save(filepath)


def generate_pdf(markdown_text: str, filepath: str):
    try:
        import markdown
        from xhtml2pdf import pisa
    except ImportError:
        raise RuntimeError("markdown or xhtml2pdf is not installed")
    
    html_content = markdown.markdown(markdown_text, extensions=['tables'])
    
    full_html = f"""
    <html>
    <head>
        <style>
            @page {{ size: a4 portrait; margin: 2cm; }}
            body {{ font-family: Helvetica, Arial, sans-serif; font-size: 12pt; line-height: 1.5; color: #333; }}
            h1 {{ color: #1a365d; border-bottom: 1px solid #ccc; padding-bottom: 5px; }}
            h2 {{ color: #2b6cb0; margin-top: 20px; }}
            h3 {{ color: #2c5282; }}
            p {{ margin-bottom: 10px; }}
            li {{ margin-bottom: 5px; }}
        </style>
    </head>
    <body>
        {html_content}
    </body>
    </html>
    """
    
    with open(filepath, "w+b") as result_file:
        pisa_status = pisa.CreatePDF(full_html, dest=result_file)
        if pisa_status.err:
            raise RuntimeError("Failed to generate PDF")


@router.post("/")
async def create_draft(request: DraftRequest):
    """
    Generate a full report based on the topic.
    Returns either a .docx or .pdf file depending on user prompt.
    
    Uses streaming generation (same as chat) to avoid GPU memory issues.
    """
    logger.info(f"Received draft request: topic='{request.topic}'")
    
    # Determine format from prompt
    req_lower = request.topic.lower()
    is_pdf = "pdf" in req_lower and "docx" not in req_lower
    
    # Build a drafting-specific prompt that wraps the user's topic
    draft_query = (
        f"Write a comprehensive, professional report on the following topic. "
        f"Include a title, introduction, detailed body sections, and conclusion. "
        f"Use markdown formatting (headers, bullet points, bold). "
        f"Topic: {request.topic}"
    )
    
    try:
        import json
        from backend.rag.chain import stream_generate_answer
        
        # Use the STREAMING generator (same code path as working chat)
        # Collect all tokens into the full text
        markdown_text = ""
        async for chunk_str in stream_generate_answer(
            query=draft_query,
            domain=request.domain if request.domain and request.domain != "all" else None
        ):
            try:
                chunk_dict = json.loads(chunk_str.strip())
                if chunk_dict.get("type") == "token":
                    markdown_text += chunk_dict.get("content", "")
            except json.JSONDecodeError:
                pass
        
        markdown_text = markdown_text.strip()
        if not markdown_text:
            raise HTTPException(status_code=500, detail="LLM returned an empty response.")
            
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Failed to generate draft: {e}")
        raise HTTPException(status_code=500, detail=str(e))
        
    # Generate file in temp directory
    tmp_dir = tempfile.gettempdir()
    file_id = str(uuid.uuid4())[:8]
    
    if is_pdf:
        filename = f"draft_report_{file_id}.pdf"
        filepath = os.path.join(tmp_dir, filename)
        generate_pdf(markdown_text, filepath)
        media_type = "application/pdf"
    else:
        filename = f"draft_report_{file_id}.docx"
        filepath = os.path.join(tmp_dir, filename)
        generate_docx(markdown_text, filepath)
        media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        
    return FileResponse(
        path=filepath,
        filename=filename,
        media_type=media_type,
        background=None
    )

